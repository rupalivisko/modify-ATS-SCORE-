from flask import Flask, request, Response
import json
import re
import pdfplumber
import os
import docx
from sentence_transformers import SentenceTransformer, util
from collections import OrderedDict
from transformers import pipeline  

app = Flask(__name__)


model = SentenceTransformer('all-MiniLM-L6-v2')
classifier = pipeline("zero-shot-classification", model="facebook/bart-large-mnli")  

def preprocess_text(text):
    text = text.lower().strip()
    text = re.sub(r"\s+", " ", text)
    return text

def extract_pdf_text(file_path):
    with pdfplumber.open(file_path) as pdf:
        return "\n".join([page.extract_text() or "" for page in pdf.pages])

def extract_word_text(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def calculate_keyword_score(resume_text):
    words = re.findall(r'\w+', resume_text.lower())
    unique_words = set(words)
    return round((len(unique_words) / len(words)) * 100, 2) if words else 0

def calculate_similarity_score(resume_text, job_text):
    if not job_text.strip():
        return 0 
    resume_embedding = model.encode(resume_text, convert_to_tensor=True)
    job_embedding = model.encode(job_text, convert_to_tensor=True)
    return round(util.pytorch_cos_sim(resume_embedding, job_embedding).item() * 100, 2)

def categorize_fit(score):
    if score >= 90:
        return "Excellent Fit"
    elif 80 <= score < 90:
        return "Very Good Fit"
    elif 70 <= score < 80:
        return "Good Fit"
    elif 60 <= score < 70:
        return "Average Fit"
    else:
        return "Not Fit"

def categorize_jd_fit(score):
    if score >= 80:
        return "Excellent Fit"
    elif 70 <= score < 80:
        return "Good Fit"
    elif 60 <= score < 70:
        return "Average Fit"
    else:
        return "Not Fit"

def check_color_alignment_fonts(file_path):
    color_used = set()
    alignments = {"left": 0, "center": 0, "right": 0}
    fonts_used = set()

    if file_path.endswith(".pdf"):
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                for char in page.chars:
                    if "non_stroking_color" in char:
                        color_used.add(char["non_stroking_color"])
                    if "fontname" in char:
                        fonts_used.add(char["fontname"])
    else:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.alignment == 0:
                alignments["left"] += 1
            elif para.alignment == 1:
                alignments["center"] += 1
            elif para.alignment == 2:
                alignments["right"] += 1
            for run in para.runs:
                if run.font and run.font.name:
                    fonts_used.add(run.font.name)

    return {
        "colors_used": list(color_used),
        "alignment_counts": alignments,
        "fonts_used": list(fonts_used)
    }

def calculate_ats_score(resume_text, color_analysis):
    words = re.findall(r'\w+', resume_text.lower())
    unique_words = set(words)
    text_score = round((len(unique_words) / len(words)) * 100, 2) if words else 0
    color_penalty = len(color_analysis["colors_used"]) * 2
    color_penalty = min(color_penalty, 10)
    total_alignments = sum(color_analysis["alignment_counts"].values())
    left_ratio = color_analysis["alignment_counts"]["left"] / total_alignments if total_alignments else 1
    alignment_penalty = (1 - left_ratio) * 10
    ats_score = text_score - color_penalty - alignment_penalty
    return max(ats_score, 0)

@app.route('/upload', methods=['POST'])
def upload_resume():
    if 'file' not in request.files:
        return Response(json.dumps({"error": "No file uploaded"}, indent=4), mimetype='application/json', status=400)

    file = request.files['file']
    if not file.filename.endswith(('.pdf', '.docx')):
        return Response(json.dumps({"error": "Only PDF or Word files are allowed"}, indent=4), mimetype='application/json', status=400)

    job_description = preprocess_text(request.form.get('job_description', ''))
    job_title = preprocess_text(request.form.get('job_title', ''))
    job_industry = preprocess_text(request.form.get('job_industry', ''))
    company_city = preprocess_text(request.form.get('company_city', ''))
    company_name = preprocess_text(request.form.get('company_name', ''))

    temp_job_description = f"Job title: {job_title}\nJob industry: {job_industry}\nCompany city: {company_city}\nCompany name: {company_name}\nJob description: {job_description}"
    temp_job_description = preprocess_text(temp_job_description)

    file_path = f"./{file.filename}"
    file.save(file_path)

    try:
        resume_text = extract_pdf_text(file_path) if file.filename.endswith('.pdf') else extract_word_text(file_path)
        resume_text = preprocess_text(resume_text)

        color_alignment_check = check_color_alignment_fonts(file_path)
        ats_score = calculate_ats_score(resume_text, color_alignment_check)
        ats_fit_status = categorize_fit(ats_score)

        similarity_scores = calculate_similarity_score(resume_text, temp_job_description) 
        jd_fit_status = categorize_jd_fit(similarity_scores)

        
        labels = ["Software Engineer", "Data Scientist", "Marketing", "Sales", "Finance", "HR"]
        classification = classifier(resume_text[:1000], labels)  
        classified_role = classification['labels'][0]  

        response = OrderedDict([
            ("ats_score", ats_score),
            ("job_description_score", similarity_scores),
            ("ats_fit_status", ats_fit_status),
            ("job_description_fit_status", jd_fit_status),
           
        ])

        return Response(json.dumps(response, indent=4, sort_keys=False), mimetype='application/json')

    except Exception as e:
        return Response(json.dumps({"error": str(e)}, indent=4), mimetype='application/json', status=500)

    finally:
        if os.path.exists(file_path):
            os.remove(file_path)

if __name__ == '__main__':
    app.run(debug=True, port=5006)
